"""Document-wide style setup + shared paragraph/table helpers for exported
exam papers (docs/export.md 版面樣式集中在樣式設定，不散落各 render 函式).

Every render function in `backend.export.renderers` builds its output
through the helpers here instead of touching `python-docx` font/paragraph
properties directly, so font/size/spacing changes happen in exactly one
place.
"""

from typing import Final, NamedTuple

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.styles.style import ParagraphStyle
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# CJK 字型選擇：.docx 只存字型「名稱」，實際字形由開啟檔案的應用程式（不是產生
# 檔案的 container）解析 —— 產生檔案的 python:3.13-slim container 本身完全不需要
# 安裝字型也能寫出正確檔案。選字依據因此是「使用者最終會在哪裡開啟、列印」而非
# container 內裝了什麼字型：本系統是台灣使用者的本機自架應用（docs/overview.md），
# 匯出檔案的用途是直接列印考卷，微軟正黑體是 Windows / Microsoft Office 繁體中文
# 版的內建字型，選它確保使用者在自己機器上開啟就是所見即所得，不必額外安裝字型。
CJK_FONT_NAME: Final[str] = "微軟正黑體"
LATIN_FONT_NAME: Final[str] = "Times New Roman"

TITLE_FONT_SIZE_PT: Final[float] = 16.0
SUBTITLE_FONT_SIZE_PT: Final[float] = 13.0
SECTION_HEADING_FONT_SIZE_PT: Final[float] = 14.0
BODY_FONT_SIZE_PT: Final[float] = 12.0
LINE_SPACING: Final[float] = 1.5

# Choice/true-false questions render this blank field next to their number
# (docs/export.md 選擇/是非題自動編號與配分欄位) — a blank for the paper's
# author to fill in by hand; the system has no per-question point-value data.
POINTS_FIELD_LABEL: Final[str] = "（配分：______分）"

# 卷首學生資訊列（docs/export.md 僅列 header_fields 勾選的欄位；順序固定為
# 班級／座號／姓名，全不勾則整列省略）。`class_` avoids shadowing the `class`
# keyword — the API-facing name (`class`) lives only in `schemas/export.py`'s
# alias, this module only ever deals with the plain Python field name.
_STUDENT_INFO_FIELD_LABELS: Final[list[tuple[str, str]]] = [
    ("class_", "班級：______________"),
    ("seat", "座號：______________"),
    ("name", "姓名：______________"),
]


class HeaderFields(NamedTuple):
    """Which 卷首 fields to print (docs/export.md 表頭選項) — `class_`/
    `seat`/`name` gate the student-info line, `score` gates the 總分 line."""

    class_: bool
    seat: bool
    name: bool
    score: bool


DEFAULT_HEADER_FIELDS: Final[HeaderFields] = HeaderFields(
    class_=True, seat=True, name=True, score=True
)


def _apply_run_font(run: Run, *, size_pt: float, bold: bool = False) -> None:
    run.font.name = LATIN_FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # `python-docx`'s `Font` has no `east_asian_name` property (its own docs
    # point users at the underlying `w:rFonts/@w:eastAsia` for this) — reached
    # via `Font.element`, a public property, so no private attribute access
    # is involved.
    r_pr = run.font.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), CJK_FONT_NAME)


def apply_base_style(document: Document) -> None:
    """Set the `Normal` style's font/size/line spacing so every paragraph
    that doesn't override it inherits the exam-paper look."""
    normal = document.styles["Normal"]
    # `document.styles["Normal"]` is statically typed as the generic
    # `BaseStyle` (styles come in several kinds); "Normal" is always a
    # paragraph style, so this narrows to the subclass that actually has
    # `.font`/`.paragraph_format` rather than widening the check away.
    if not isinstance(normal, ParagraphStyle):
        raise TypeError(f"expected the 'Normal' style to be a paragraph style, got {type(normal)}")
    normal.font.name = LATIN_FONT_NAME
    normal.font.size = Pt(BODY_FONT_SIZE_PT)
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), CJK_FONT_NAME)
    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = LINE_SPACING


def add_title(document: Document, text: str) -> Paragraph:
    """Centered, bold, larger-font document title (e.g. 題目卷/答案卷)."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _apply_run_font(run, size_pt=TITLE_FONT_SIZE_PT, bold=True)
    return paragraph


def add_subtitle(document: Document, text: str) -> Paragraph:
    """Centered, bold, smaller-than-title line under the exam title — used
    for the 題目卷/答案卷 label so the exam's own title (docs/export.md
    卷首：考卷標題) stays the visually dominant heading."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _apply_run_font(run, size_pt=SUBTITLE_FONT_SIZE_PT, bold=True)
    return paragraph


def add_student_info_line(document: Document, header_fields: HeaderFields) -> Paragraph | None:
    """卷首學生資訊列：僅印 `header_fields` 勾選的欄位（班級／座號／姓名子
    集，固定順序，既有間距）；全不勾時整列省略，回傳 `None`（docs/export.md）。
    """
    parts = [
        label
        for field_name, label in _STUDENT_INFO_FIELD_LABELS
        if getattr(header_fields, field_name)
    ]
    if not parts:
        return None
    paragraph = document.add_paragraph()
    run = paragraph.add_run("　".join(parts))
    _apply_run_font(run, size_pt=BODY_FONT_SIZE_PT)
    return paragraph


def add_total_score_line(document: Document, total: int) -> Paragraph:
    """卷首總分 (docs/export.md 有配分時印總分) — only ever called when at
    least one type in the export's `points` mapping was actually assigned."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"總分：{total} 分")
    _apply_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=True)
    return paragraph


def add_section_heading(document: Document, text: str) -> Paragraph:
    """One 分節 heading (e.g. 「一、選擇題（每題 2 分）」), left-aligned and
    bold so it stands out from the numbered stems under it without competing
    with the exam title's centered/larger styling."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    run = paragraph.add_run(text)
    _apply_run_font(run, size_pt=SECTION_HEADING_FONT_SIZE_PT, bold=True)
    return paragraph


def add_numbered_stem(
    document: Document,
    number: int,
    text: str,
    *,
    points_field: bool = False,
    points_suffix: int | None = None,
) -> Paragraph:
    """`{number}. {text}`, optionally followed by the 配分 blank field
    (`points_field`) or preceded by a 「（X 分）」 resolved-points suffix
    right after the number (`points_suffix`, docs/export.md 節內配分不一致
    時各題題號後印「（X 分）」) — the two are mutually exclusive by
    construction (a question either still has the old hand-fill blank, or it
    already has a resolved points value, never both)."""
    paragraph = document.add_paragraph()
    prefix = f"{number}"
    if points_suffix is not None:
        prefix = f"{prefix}（{points_suffix} 分）"
    prefix = f"{prefix}. "
    if points_field:
        prefix = f"{prefix}{POINTS_FIELD_LABEL}　"
    run = paragraph.add_run(prefix + text)
    _apply_run_font(run, size_pt=BODY_FONT_SIZE_PT)
    return paragraph


def add_body_paragraph(document: Document, text: str, *, indent: bool = False) -> Paragraph:
    """A plain paragraph (option line, answer line, explanation, ...);
    `indent=True` for lines that belong under a numbered stem (options,
    bullet points)."""
    paragraph = document.add_paragraph()
    if indent:
        paragraph.paragraph_format.left_indent = Pt(24)
    run = paragraph.add_run(text)
    _apply_run_font(run, size_pt=BODY_FONT_SIZE_PT)
    return paragraph


def add_answer_table(document: Document, headers: list[str], rows: list[list[str]]) -> Table:
    """A bordered `len(headers)`-column table — used for the comparison
    question's 異同表 (面向 x A x B) on the answer sheet."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers, strict=True):
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        _apply_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values, strict=True):
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            _apply_run_font(run, size_pt=BODY_FONT_SIZE_PT)
    return table
