"""`export_docx` job handler tests (.rule 反偷懶規則 — Word 匯出屬高風險邏輯，須有測試).

Real Postgres + real `python-docx` rendering (no network/LLM involved) —
covers approved-only validation (missing/not-approved ids fail the whole
job, per docs/question-bank.md 只有 approved 題目...匯出的選題範圍), files
actually written under a throwaway `DATA_DIR`, the inserted `exports` row,
and per-question progress reporting.
"""

from pathlib import Path

import pytest
from docx import Document as open_docx
from docx.document import Document
from factories import create_job
from sqlalchemy import select

from backend.core.config import get_settings
from backend.db.session import AsyncSessionLocal
from backend.jobs.worker import claim_job, run_claimed_job
from backend.models.export import Export
from backend.models.job import Job
from backend.models.question import Question

SINGLE_CHOICE_PAYLOAD: dict[str, object] = {
    "stem": "光合作用發生在細胞的哪個構造？",
    "options": ["粒線體", "葉綠體", "細胞核", "核糖體"],
    "answer_index": 1,
    "explanation": None,
}

TRUE_FALSE_PAYLOAD: dict[str, object] = {
    "stem": "光合作用會釋放氧氣。",
    "answer": True,
    "explanation": None,
}


async def _make_question(
    *,
    question_type: str = "single_choice",
    status: str = "approved",
    payload: dict[str, object] | None = None,
) -> int:
    async with AsyncSessionLocal() as session:
        question = Question(
            type=question_type,
            status=status,
            payload=payload or SINGLE_CHOICE_PAYLOAD,
            source_chunk_ids=[],
        )
        session.add(question)
        await session.commit()
        await session.refresh(question)
        return question.id


DEFAULT_TITLE = "測試考卷"


async def _run_job(payload: dict[str, object]) -> int:
    """`title` defaults to `DEFAULT_TITLE` for tests that don't care about
    it — only the title/points-specific tests below override it."""
    full_payload = {"title": DEFAULT_TITLE, **payload}
    job_id = await create_job("export_docx", payload=full_payload)
    async with AsyncSessionLocal() as session:
        claimed = await claim_job(session)
        assert claimed is not None
        assert claimed.id == job_id
    await run_claimed_job(AsyncSessionLocal, job_id)
    return job_id


async def _get_job(job_id: int) -> Job:
    async with AsyncSessionLocal() as session:
        job = await session.get(Job, job_id)
        assert job is not None
        return job


async def test_rejects_not_approved_question_id_and_names_it(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")
    draft_id = await _make_question(status="draft")

    job_id = await _run_job({"question_ids": [approved_id, draft_id], "paper_size": "A4"})
    job = await _get_job(job_id)

    assert job.status == "failed"
    assert job.error is not None
    assert str(draft_id) in job.error
    assert "not approved" in job.error

    async with AsyncSessionLocal() as session:
        exports = (await session.execute(select(Export))).scalars().all()
    assert exports == []  # a rejected job creates no export row

    get_settings.cache_clear()


async def test_rejects_missing_question_id_and_names_it(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")
    missing_id = approved_id + 999999

    job_id = await _run_job({"question_ids": [approved_id, missing_id], "paper_size": "A4"})
    job = await _get_job(job_id)

    assert job.status == "failed"
    assert job.error is not None
    assert str(missing_id) in job.error
    assert "not found" in job.error

    get_settings.cache_clear()


async def test_unsupported_paper_size_fails_before_touching_the_database(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")

    job_id = await _run_job({"question_ids": [approved_id], "paper_size": "Letter"})
    job = await _get_job(job_id)

    assert job.status == "failed"
    assert job.error is not None
    assert "Letter" in job.error

    async with AsyncSessionLocal() as session:
        exports = (await session.execute(select(Export))).scalars().all()
    assert exports == []

    get_settings.cache_clear()


async def test_writes_both_files_and_correct_exports_row(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    single_choice_id = await _make_question(status="approved")
    true_false_id = await _make_question(
        status="approved", question_type="true_false", payload=TRUE_FALSE_PAYLOAD
    )

    job_id = await _run_job(
        {"question_ids": [single_choice_id, true_false_id], "paper_size": "B4"}
    )
    job = await _get_job(job_id)

    assert job.status == "done"
    assert job.error is None
    assert job.progress == "2/2"

    async with AsyncSessionLocal() as session:
        exports = (await session.execute(select(Export))).scalars().all()
    assert len(exports) == 1
    export = exports[0]
    assert export.paper_size == "B4"
    assert export.question_ids == [single_choice_id, true_false_id]
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    questions_path = Path(export.docx_path)
    answers_path = Path(export.answer_docx_path)
    assert questions_path.exists()
    assert answers_path.exists()
    assert questions_path.is_relative_to(tmp_path / "exports")

    question_doc = open_docx(str(questions_path))
    answer_doc = open_docx(str(answers_path))
    stem_paragraphs = [
        paragraph.text for paragraph in question_doc.paragraphs if paragraph.text[:1].isdigit()
    ]
    assert len(stem_paragraphs) == 2
    assert any("答案：○（正確）" in paragraph.text for paragraph in answer_doc.paragraphs)

    get_settings.cache_clear()


async def test_progress_follows_generate_questions_n_of_total_convention(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Final `jobs.progress` after a 3-question export must read "3/3", the
    same "N/total" convention `generate_questions` uses — the intermediate
    per-question commits (`ctx.set_progress` called once per rendered
    question in `backend.export.job.export_docx`) are what make that number
    correct; a batch call that only reported 0%/100% would also end at
    "3/3" by coincidence, so this is paired with the renderer/job unit tests
    that check each question actually got rendered, not just counted."""
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    question_ids = [await _make_question(status="approved") for _ in range(3)]

    job_id = await _run_job({"question_ids": question_ids, "paper_size": "A4"})
    job = await _get_job(job_id)

    assert job.status == "done"
    assert job.progress == "3/3"

    get_settings.cache_clear()


async def test_duplicate_question_ids_are_deduplicated_and_numbered_once(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    question_id = await _make_question(status="approved")

    job_id = await _run_job(
        {"question_ids": [question_id, question_id], "paper_size": "A4"}
    )
    job = await _get_job(job_id)

    assert job.status == "done"
    assert job.progress == "1/1"

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.question_ids == [question_id]

    get_settings.cache_clear()


# ---------------------------------------------------------------------------
# title / sections / points (docs/export.md 卷面結構)
# ---------------------------------------------------------------------------


def _all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


async def test_missing_title_fails_the_job(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")

    job_id = await create_job(
        "export_docx", payload={"question_ids": [approved_id], "paper_size": "A4"}
    )
    async with AsyncSessionLocal() as session:
        claimed = await claim_job(session)
        assert claimed is not None
        assert claimed.id == job_id
    await run_claimed_job(AsyncSessionLocal, job_id)
    job = await _get_job(job_id)

    assert job.status == "failed"
    assert job.error is not None
    assert "title" in job.error

    get_settings.cache_clear()


async def test_exports_row_and_docx_header_carry_the_title(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    question_id = await _make_question(status="approved")

    job_id = await _run_job(
        {"question_ids": [question_id], "paper_size": "A4", "title": "第一次段考"}
    )
    job = await _get_job(job_id)
    assert job.status == "done"

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.title == "第一次段考"
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    question_doc = open_docx(str(Path(export.docx_path)))
    answer_doc = open_docx(str(Path(export.answer_docx_path)))
    for document in (question_doc, answer_doc):
        text = _all_text(document)
        assert "第一次段考" in text
        assert "班級" in text and "座號" in text and "姓名" in text

    get_settings.cache_clear()


async def test_sections_grouped_headed_and_renumbered_with_points_and_total_score(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Two single_choice + one true_false, requested in an interleaved
    order, must render as two sections (選擇題 first per the fixed order,
    regardless of request order), each renumbered from 1, with the
    single_choice heading carrying its assigned 配分 and the true_false
    section keeping its old per-question 配分 blank (not assigned points)."""
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    true_false_id = await _make_question(
        status="approved", question_type="true_false", payload=TRUE_FALSE_PAYLOAD
    )
    single_choice_a = await _make_question(status="approved")
    single_choice_b = await _make_question(status="approved")

    job_id = await _run_job(
        {
            "question_ids": [true_false_id, single_choice_a, single_choice_b],
            "paper_size": "A4",
            "title": "小考",
            "points": {"single_choice": 5},
        }
    )
    job = await _get_job(job_id)
    assert job.status == "done", job.error

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    for path_str in (export.docx_path, export.answer_docx_path):
        document = open_docx(str(Path(path_str)))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]

        assert "一、選擇題（每題 5 分）" in paragraphs
        assert "二、是非題" in paragraphs  # no points assigned -> no "每題 X 分"

        section_index = paragraphs.index("一、選擇題（每題 5 分）")
        true_false_index = paragraphs.index("二、是非題")
        single_choice_numbers = [
            paragraph[:1]
            for paragraph in paragraphs[section_index:true_false_index]
            if paragraph[:1].isdigit()
        ]
        assert single_choice_numbers == ["1", "2"]  # renumbered from 1, not 2/3

        true_false_stem = next(
            paragraph for paragraph in paragraphs[true_false_index:] if paragraph[:1].isdigit()
        )
        assert true_false_stem.startswith("1. ")  # true_false section also restarts at 1
        assert "配分" in true_false_stem  # kept: true_false has no assigned points

        assert "總分：10 分" in paragraphs  # 2 single_choice * 5 分; true_false contributes 0

    get_settings.cache_clear()


# ---------------------------------------------------------------------------
# question_points (逐題覆寫) / header_fields (docs/export.md 卷面結構)
# ---------------------------------------------------------------------------


async def test_rejects_question_points_key_outside_question_ids(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")
    outside_id = approved_id + 999999

    job_id = await _run_job(
        {
            "question_ids": [approved_id],
            "paper_size": "A4",
            "question_points": {str(outside_id): 5},
        }
    )
    job = await _get_job(job_id)

    assert job.status == "failed"
    assert job.error is not None
    assert "question_points" in job.error
    assert str(outside_id) in job.error

    get_settings.cache_clear()


async def test_rejects_non_positive_question_points_value(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")

    job_id = await _run_job(
        {
            "question_ids": [approved_id],
            "paper_size": "A4",
            "question_points": {str(approved_id): 0},
        }
    )
    job = await _get_job(job_id)

    assert job.status == "failed"
    assert job.error is not None
    assert "positive" in job.error

    get_settings.cache_clear()


async def test_question_points_override_wins_over_type_points_in_total_and_heading(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Two single_choice questions with a 5-point type default, one of them
    overridden to 9 -- the section becomes non-uniform (5 vs 9) so the
    heading drops "每題 X 分" and each question gets its own suffix instead,
    and the total sums the *resolved* per-question values (9 + 5), not
    2 * the type default."""
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    single_choice_a = await _make_question(status="approved")
    single_choice_b = await _make_question(status="approved")

    job_id = await _run_job(
        {
            "question_ids": [single_choice_a, single_choice_b],
            "paper_size": "A4",
            "points": {"single_choice": 5},
            "question_points": {str(single_choice_a): 9},
        }
    )
    job = await _get_job(job_id)
    assert job.status == "done", job.error

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    for path_str in (export.docx_path, export.answer_docx_path):
        document = open_docx(str(Path(path_str)))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]

        assert "一、選擇題" in paragraphs  # not uniform -> no "（每題 X 分）"
        assert "一、選擇題（每題 5 分）" not in paragraphs
        stem_lines = [p for p in paragraphs if p[:1].isdigit()]
        assert any(line.startswith("1（9 分）.") for line in stem_lines)
        assert any(line.startswith("2（5 分）.") for line in stem_lines)
        assert "配分：" not in "\n".join(stem_lines)  # resolved points -> no hand-fill blank

        assert "總分：14 分" in paragraphs  # 9 (override) + 5 (type default)

    get_settings.cache_clear()


async def test_uniform_section_prints_per_type_heading_without_per_question_suffix(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    single_choice_a = await _make_question(status="approved")
    single_choice_b = await _make_question(status="approved")

    job_id = await _run_job(
        {
            "question_ids": [single_choice_a, single_choice_b],
            "paper_size": "A4",
            "points": {"single_choice": 5},
        }
    )
    job = await _get_job(job_id)
    assert job.status == "done", job.error

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.docx_path is not None

    document = open_docx(str(Path(export.docx_path)))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert "一、選擇題（每題 5 分）" in paragraphs
    stem_lines = [p for p in paragraphs if p[:1].isdigit()]
    assert stem_lines[0].startswith("1. ")  # uniform -> plain number, no suffix
    assert stem_lines[1].startswith("2. ")
    assert "（" not in stem_lines[0].split(".", 1)[0]

    get_settings.cache_clear()


async def test_questions_without_any_points_keep_the_blank_field(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")

    job_id = await _run_job({"question_ids": [approved_id], "paper_size": "A4"})
    job = await _get_job(job_id)
    assert job.status == "done", job.error

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    for path_str in (export.docx_path, export.answer_docx_path):
        document = open_docx(str(Path(path_str)))
        text = _all_text(document)
        assert "配分：______分" in text
        assert "總分" not in text  # nothing scored -> no total line at all

    get_settings.cache_clear()


async def test_header_line_shows_only_the_checked_field_subset(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")

    job_id = await _run_job(
        {
            "question_ids": [approved_id],
            "paper_size": "A4",
            "header_fields": {"class": False, "seat": True, "name": True, "score": True},
        }
    )
    job = await _get_job(job_id)
    assert job.status == "done", job.error

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    for path_str in (export.docx_path, export.answer_docx_path):
        text = _all_text(open_docx(str(Path(path_str))))
        assert "班級" not in text
        assert "座號" in text
        assert "姓名" in text

    get_settings.cache_clear()


async def test_header_line_omitted_entirely_when_no_field_checked(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")

    job_id = await _run_job(
        {
            "question_ids": [approved_id],
            "paper_size": "A4",
            "header_fields": {"class": False, "seat": False, "name": False, "score": True},
        }
    )
    job = await _get_job(job_id)
    assert job.status == "done", job.error

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    for path_str in (export.docx_path, export.answer_docx_path):
        text = _all_text(open_docx(str(Path(path_str))))
        assert "班級" not in text
        assert "座號" not in text
        assert "姓名" not in text

    get_settings.cache_clear()


async def test_total_score_hidden_when_score_field_is_off_even_with_points(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()

    approved_id = await _make_question(status="approved")

    job_id = await _run_job(
        {
            "question_ids": [approved_id],
            "paper_size": "A4",
            "points": {"single_choice": 5},
            "header_fields": {"class": True, "seat": True, "name": True, "score": False},
        }
    )
    job = await _get_job(job_id)
    assert job.status == "done", job.error

    async with AsyncSessionLocal() as session:
        export = (await session.execute(select(Export))).scalars().one()
    assert export.docx_path is not None
    assert export.answer_docx_path is not None

    for path_str in (export.docx_path, export.answer_docx_path):
        text = _all_text(open_docx(str(Path(path_str))))
        assert "總分" not in text

    get_settings.cache_clear()
