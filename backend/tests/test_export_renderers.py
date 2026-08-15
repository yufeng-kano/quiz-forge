"""Renderer tests (.rule 反偷懶規則 — Word 匯出屬高風險邏輯，須有測試).

No network, no LLM cost: build one sample Pydantic model per question type
directly, render both papers through `ExamPaperBuilder`, then re-open the
saved `.docx` with `python-docx` and assert on real paragraph/table content
— never a bare "file exists"/"is not None".
"""

from pathlib import Path

import pytest
from docx import Document as open_docx
from docx.document import Document

from backend.export.builder import ExamPaperBuilder
from backend.export.paper import PAGE_MARGIN_MM, PAPER_SIZES_MM
from backend.questions.schemas import (
    AnalogyQuestion,
    ComparisonDifference,
    ComparisonModelAnswer,
    ComparisonQuestion,
    FillBlankQuestion,
    QuestionModel,
    ShortAnswerQuestion,
    SingleChoiceQuestion,
    TrueFalseQuestion,
)

COMPARISON = ComparisonQuestion(
    stem="試比較光合作用與呼吸作用之異同。",
    subject_a="光合作用",
    subject_b="呼吸作用",
    aspects=["場所"],
    model_answer=ComparisonModelAnswer(
        similarities=["皆為細胞內能量代謝反應"],
        differences=[ComparisonDifference(aspect="場所", a="葉綠體", b="粒線體")],
    ),
)

ANALOGY_CHOICE = AnalogyQuestion(
    a="筆",
    b="寫字",
    c="剪刀",
    answer="剪裁",
    options=["剪裁", "縫紉", "烹飪", "測量"],
    explanation="工具之於其功能",
)

ANALOGY_FILL_IN = AnalogyQuestion(
    a="醫生", b="醫院", c="老師", answer="學校", options=None, explanation=None
)

SINGLE_CHOICE = SingleChoiceQuestion(
    stem="光合作用發生在細胞的哪個構造？",
    options=["粒線體", "葉綠體", "細胞核", "核糖體"],
    answer_index=1,
    explanation="葉綠體含葉綠素",
)

TRUE_FALSE = TrueFalseQuestion(
    stem="光合作用會釋放氧氣。", answer=True, explanation="產物之一為氧氣"
)

FILL_BLANK = FillBlankQuestion(stem="水的化學式為 ____，由氫與 ____ 組成。", answers=["H2O", "氧"])

SHORT_ANSWER = ShortAnswerQuestion(
    stem="請說明光合作用的功能。",
    model_answer="將光能轉換成化學能。",
    key_points=["光能轉化學能", "產生葡萄糖"],
)


def _all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _build_and_reopen(
    tmp_path: Path, paper_size: str, questions: list[QuestionModel]
) -> tuple[Document, Document]:
    builder = ExamPaperBuilder(paper_size)
    for number, question in enumerate(questions, start=1):
        builder.render_question(number, question)
    questions_path = tmp_path / "questions.docx"
    answers_path = tmp_path / "answers.docx"
    builder.save(questions_path, answers_path)
    return open_docx(str(questions_path)), open_docx(str(answers_path))


def test_comparison_renders_stem_and_answer_sheet_similarity_table(tmp_path: Path) -> None:
    question_doc, answer_doc = _build_and_reopen(tmp_path, "A4", [COMPARISON])

    assert "試比較光合作用與呼吸作用之異同。" in _all_text(question_doc)
    assert not question_doc.tables, "題目卷不應含異同表"

    assert len(answer_doc.tables) == 1
    table = answer_doc.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == ["面向", "光合作用", "呼吸作用"]
    assert [cell.text for cell in table.rows[1].cells] == ["場所", "葉綠體", "粒線體"]
    assert "皆為細胞內能量代謝反應" in _all_text(answer_doc)


def test_analogy_stem_composed_from_slots_renders_choice_form(tmp_path: Path) -> None:
    question_doc, answer_doc = _build_and_reopen(tmp_path, "A4", [ANALOGY_CHOICE])

    question_text = _all_text(question_doc)
    assert "筆 之於 寫字，猶如 剪刀 之於＿＿" in question_text
    assert "(A) 剪裁" in question_text
    assert "答案" not in question_text

    answer_text = _all_text(answer_doc)
    assert "答案：剪裁" in answer_text
    assert "工具之於其功能" in answer_text


def test_analogy_options_none_renders_fill_in_form_without_option_lines(tmp_path: Path) -> None:
    question_doc, _ = _build_and_reopen(tmp_path, "A4", [ANALOGY_FILL_IN])

    question_text = _all_text(question_doc)
    assert "醫生 之於 醫院，猶如 老師 之於＿＿" in question_text
    assert "(A)" not in question_text


def test_single_choice_options_lettered_answer_only_on_answer_sheet(tmp_path: Path) -> None:
    question_doc, answer_doc = _build_and_reopen(tmp_path, "A4", [SINGLE_CHOICE])

    question_text = _all_text(question_doc)
    assert "光合作用發生在細胞的哪個構造？" in question_text
    assert "(A) 粒線體" in question_text
    assert "(B) 葉綠體" in question_text
    assert "配分" in question_text  # 選擇題自動配分欄位
    assert "答案" not in question_text

    answer_text = _all_text(answer_doc)
    assert "答案：(B) 葉綠體" in answer_text
    assert "葉綠體含葉綠素" in answer_text


def test_true_false_gets_points_field_and_answer_symbol(tmp_path: Path) -> None:
    question_doc, answer_doc = _build_and_reopen(tmp_path, "A4", [TRUE_FALSE])

    assert "配分" in _all_text(question_doc)
    assert "答案：○（正確）" in _all_text(answer_doc)
    assert "產物之一為氧氣" in _all_text(answer_doc)


def test_fill_blank_keeps_blank_markers_and_answers_only_on_answer_sheet(tmp_path: Path) -> None:
    question_doc, answer_doc = _build_and_reopen(tmp_path, "A4", [FILL_BLANK])

    question_text = _all_text(question_doc)
    assert "水的化學式為 ____，由氫與 ____ 組成。" in question_text
    assert "H2O" not in question_text

    answer_text = _all_text(answer_doc)
    assert "1. H2O" in answer_text
    assert "2. 氧" in answer_text


def test_short_answer_model_answer_and_key_points_only_on_answer_sheet(tmp_path: Path) -> None:
    question_doc, answer_doc = _build_and_reopen(tmp_path, "A4", [SHORT_ANSWER])

    assert "參考答案" not in _all_text(question_doc)

    answer_text = _all_text(answer_doc)
    assert "參考答案：將光能轉換成化學能。" in answer_text
    assert "光能轉化學能" in answer_text
    assert "產生葡萄糖" in answer_text


def test_numbering_is_sequential_across_mixed_question_types(tmp_path: Path) -> None:
    questions: list[QuestionModel] = [
        COMPARISON,
        ANALOGY_CHOICE,
        SINGLE_CHOICE,
        TRUE_FALSE,
        FILL_BLANK,
        SHORT_ANSWER,
    ]
    question_doc, answer_doc = _build_and_reopen(tmp_path, "A4", questions)

    for document in (question_doc, answer_doc):
        stem_paragraphs = [
            paragraph.text for paragraph in document.paragraphs if paragraph.text[:1].isdigit()
        ]
        numbers = [int(text.split(".", 1)[0]) for text in stem_paragraphs]
        assert numbers == [1, 2, 3, 4, 5, 6]


@pytest.mark.parametrize("paper_size", sorted(PAPER_SIZES_MM))
def test_page_dimensions_match_mm_constants_for_every_supported_size(
    tmp_path: Path, paper_size: str
) -> None:
    question_doc, answer_doc = _build_and_reopen(tmp_path, paper_size, [SINGLE_CHOICE])
    width_mm, height_mm = PAPER_SIZES_MM[paper_size]

    for document in (question_doc, answer_doc):
        section = document.sections[0]
        page_width = section.page_width
        page_height = section.page_height
        top_margin = section.top_margin
        bottom_margin = section.bottom_margin
        left_margin = section.left_margin
        right_margin = section.right_margin
        assert page_width is not None
        assert page_height is not None
        assert top_margin is not None
        assert bottom_margin is not None
        assert left_margin is not None
        assert right_margin is not None
        assert abs(page_width.mm - width_mm) < 0.1
        assert abs(page_height.mm - height_mm) < 0.1
        assert abs(top_margin.mm - PAGE_MARGIN_MM) < 0.1
        assert abs(bottom_margin.mm - PAGE_MARGIN_MM) < 0.1
        assert abs(left_margin.mm - PAGE_MARGIN_MM) < 0.1
        assert abs(right_margin.mm - PAGE_MARGIN_MM) < 0.1
