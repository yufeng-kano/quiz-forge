"""`/v1/exports` through the real HTTP app.

The `client` fixture disables the job worker pool, so `POST /v1/exports`
here only exercises row/job creation and paper-size validation (the
`export_docx` handler itself — approved-only validation, real file writes,
`exports` row correctness — is covered against real Postgres + real
`python-docx` in `test_export_job.py`)."""

from pathlib import Path
from urllib.parse import unquote

import pytest
from docx import Document as new_docx
from fastapi.testclient import TestClient

from backend.db.session import AsyncSessionLocal
from backend.models.export import Export
from backend.models.job import Job
from backend.models.question import Question

SINGLE_CHOICE_PAYLOAD: dict[str, object] = {
    "stem": "...",
    "options": ["a", "b", "c", "d"],
    "answer_index": 0,
    "explanation": None,
}


async def _make_question(*, status: str = "approved") -> int:
    async with AsyncSessionLocal() as session:
        question = Question(
            type="single_choice", status=status, payload=SINGLE_CHOICE_PAYLOAD, source_chunk_ids=[]
        )
        session.add(question)
        await session.commit()
        await session.refresh(question)
        return question.id


async def _make_export(
    *,
    title: str = "測試考卷",
    paper_size: str = "A4",
    question_ids: list[int] | None = None,
    docx_path: str | None = None,
    answer_docx_path: str | None = None,
) -> int:
    async with AsyncSessionLocal() as session:
        export = Export(
            title=title,
            paper_size=paper_size,
            question_ids=question_ids or [],
            docx_path=docx_path,
            answer_docx_path=answer_docx_path,
        )
        session.add(export)
        await session.commit()
        await session.refresh(export)
        return export.id


# ---------------------------------------------------------------------------
# POST /v1/exports
# ---------------------------------------------------------------------------


async def test_create_export_job_enqueues_export_docx_job(client: TestClient) -> None:
    question_id = await _make_question()

    response = client.post(
        "/v1/exports",
        json={"question_ids": [question_id], "paper_size": "B4", "title": "第一次段考"},
    )

    assert response.status_code == 201
    body = response.json()
    assert isinstance(body["job_id"], int)

    async with AsyncSessionLocal() as session:
        job = await session.get(Job, body["job_id"])
        assert job is not None
        assert job.kind == "export_docx"
        assert job.status == "pending"
        assert job.payload == {
            "question_ids": [question_id],
            "paper_size": "B4",
            "title": "第一次段考",
            "points": None,
            "question_points": None,
            "header_fields": {"class": True, "seat": True, "name": True, "score": True},
        }


async def test_create_export_job_forwards_points(client: TestClient) -> None:
    question_id = await _make_question()

    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [question_id],
            "paper_size": "A4",
            "title": "小考",
            "points": {"single_choice": 5, "true_false": 2},
        },
    )

    assert response.status_code == 201
    async with AsyncSessionLocal() as session:
        job = await session.get(Job, response.json()["job_id"])
        assert job is not None
        assert job.payload["points"] == {"single_choice": 5, "true_false": 2}


def test_create_export_job_rejects_unsupported_paper_size(client: TestClient) -> None:
    response = client.post(
        "/v1/exports", json={"question_ids": [1], "paper_size": "Letter", "title": "考卷"}
    )
    assert response.status_code == 422


def test_create_export_job_rejects_empty_question_ids(client: TestClient) -> None:
    response = client.post(
        "/v1/exports", json={"question_ids": [], "paper_size": "A4", "title": "考卷"}
    )
    assert response.status_code == 422


def test_create_export_job_rejects_missing_title(client: TestClient) -> None:
    response = client.post("/v1/exports", json={"question_ids": [1], "paper_size": "A4"})
    assert response.status_code == 422


def test_create_export_job_rejects_blank_title(client: TestClient) -> None:
    response = client.post(
        "/v1/exports", json={"question_ids": [1], "paper_size": "A4", "title": "   "}
    )
    assert response.status_code == 422


def test_create_export_job_rejects_unknown_points_type(client: TestClient) -> None:
    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [1],
            "paper_size": "A4",
            "title": "考卷",
            "points": {"essay": 5},
        },
    )
    assert response.status_code == 422


def test_create_export_job_rejects_non_positive_points_value(client: TestClient) -> None:
    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [1],
            "paper_size": "A4",
            "title": "考卷",
            "points": {"single_choice": 0},
        },
    )
    assert response.status_code == 422


# ---------------------------------------------------------------------------
# question_points (逐題覆寫)
# ---------------------------------------------------------------------------


async def test_create_export_job_forwards_question_points(client: TestClient) -> None:
    question_id = await _make_question()

    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [question_id],
            "paper_size": "A4",
            "title": "小考",
            "question_points": {str(question_id): 8},
        },
    )

    assert response.status_code == 201
    async with AsyncSessionLocal() as session:
        job = await session.get(Job, response.json()["job_id"])
        assert job is not None
        assert job.payload["question_points"] == {str(question_id): 8}


def test_create_export_job_rejects_question_points_key_outside_question_ids(
    client: TestClient,
) -> None:
    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [1],
            "paper_size": "A4",
            "title": "考卷",
            "question_points": {"2": 5},
        },
    )
    assert response.status_code == 422
    assert "question_points" in response.text


def test_create_export_job_rejects_non_positive_question_points_value(
    client: TestClient,
) -> None:
    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [1],
            "paper_size": "A4",
            "title": "考卷",
            "question_points": {"1": 0},
        },
    )
    assert response.status_code == 422


async def test_create_export_job_question_points_override_wins_over_points(
    client: TestClient,
) -> None:
    question_id = await _make_question()

    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [question_id],
            "paper_size": "A4",
            "title": "小考",
            "points": {"single_choice": 5},
            "question_points": {str(question_id): 9},
        },
    )

    assert response.status_code == 201
    async with AsyncSessionLocal() as session:
        job = await session.get(Job, response.json()["job_id"])
        assert job is not None
        assert job.payload["points"] == {"single_choice": 5}
        assert job.payload["question_points"] == {str(question_id): 9}


# ---------------------------------------------------------------------------
# header_fields (卷首欄位開關)
# ---------------------------------------------------------------------------


async def test_create_export_job_defaults_header_fields_to_all_true(client: TestClient) -> None:
    question_id = await _make_question()

    response = client.post(
        "/v1/exports",
        json={"question_ids": [question_id], "paper_size": "A4", "title": "小考"},
    )

    assert response.status_code == 201
    async with AsyncSessionLocal() as session:
        job = await session.get(Job, response.json()["job_id"])
        assert job is not None
        assert job.payload["header_fields"] == {
            "class": True,
            "seat": True,
            "name": True,
            "score": True,
        }


async def test_create_export_job_forwards_partial_header_fields_using_class_alias(
    client: TestClient,
) -> None:
    question_id = await _make_question()

    response = client.post(
        "/v1/exports",
        json={
            "question_ids": [question_id],
            "paper_size": "A4",
            "title": "小考",
            "header_fields": {"class": False, "score": False},
        },
    )

    assert response.status_code == 201
    async with AsyncSessionLocal() as session:
        job = await session.get(Job, response.json()["job_id"])
        assert job is not None
        assert job.payload["header_fields"] == {
            "class": False,
            "seat": True,
            "name": True,
            "score": False,
        }


@pytest.mark.parametrize("paper_size", ["A4", "B4", "B3"])
def test_create_export_job_accepts_every_supported_paper_size(
    client: TestClient, paper_size: str
) -> None:
    response = client.post(
        "/v1/exports", json={"question_ids": [1], "paper_size": paper_size, "title": "考卷"}
    )
    assert response.status_code == 201


# ---------------------------------------------------------------------------
# GET /v1/exports
# ---------------------------------------------------------------------------


async def test_list_exports_reports_question_count_and_paper_size(client: TestClient) -> None:
    q1 = await _make_question()
    q2 = await _make_question()
    export_id = await _make_export(title="第一次段考", paper_size="B4", question_ids=[q1, q2])

    response = client.get("/v1/exports")

    assert response.status_code == 200
    item = next(item for item in response.json() if item["id"] == export_id)
    assert item["title"] == "第一次段考"
    assert item["paper_size"] == "B4"
    assert item["question_count"] == 2
    assert item["questions_available"] is False
    assert item["answers_available"] is False


async def test_list_exports_reports_file_availability_when_files_exist(
    client: TestClient, tmp_path: Path
) -> None:
    questions_path = tmp_path / "q.docx"
    answers_path = tmp_path / "a.docx"
    new_docx().save(str(questions_path))
    new_docx().save(str(answers_path))
    export_id = await _make_export(
        docx_path=str(questions_path), answer_docx_path=str(answers_path)
    )

    response = client.get("/v1/exports")

    item = next(item for item in response.json() if item["id"] == export_id)
    assert item["questions_available"] is True
    assert item["answers_available"] is True


async def test_list_exports_newest_first(client: TestClient) -> None:
    first_id = await _make_export()
    second_id = await _make_export()

    response = client.get("/v1/exports")

    ids = [item["id"] for item in response.json()]
    assert ids.index(second_id) < ids.index(first_id)


# ---------------------------------------------------------------------------
# GET /v1/exports/{id}/questions.docx, /answers.docx
# ---------------------------------------------------------------------------


def test_download_questions_docx_404_for_missing_export(client: TestClient) -> None:
    response = client.get("/v1/exports/999999999/questions.docx")
    assert response.status_code == 404


async def test_download_questions_docx_404_when_not_yet_rendered(client: TestClient) -> None:
    export_id = await _make_export(docx_path=None)
    response = client.get(f"/v1/exports/{export_id}/questions.docx")
    assert response.status_code == 404


async def test_download_questions_docx_404_when_file_missing_on_disk(
    client: TestClient, tmp_path: Path
) -> None:
    export_id = await _make_export(docx_path=str(tmp_path / "gone.docx"))
    response = client.get(f"/v1/exports/{export_id}/questions.docx")
    assert response.status_code == 404


async def test_download_questions_docx_returns_valid_docx_with_content_type_and_filename(
    client: TestClient, tmp_path: Path
) -> None:
    questions_path = tmp_path / "questions.docx"
    document = new_docx()
    document.add_paragraph("1. 測試題幹")
    document.save(str(questions_path))
    export_id = await _make_export(docx_path=str(questions_path))

    response = client.get(f"/v1/exports/{export_id}/questions.docx")

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "題目卷" in unquote(response.headers["content-disposition"])

    downloaded_path = tmp_path / "downloaded-questions.docx"
    downloaded_path.write_bytes(response.content)
    reopened = new_docx(str(downloaded_path))
    assert reopened.paragraphs[0].text == "1. 測試題幹"


async def test_download_answers_docx_returns_valid_docx(client: TestClient, tmp_path: Path) -> None:
    answers_path = tmp_path / "answers.docx"
    document = new_docx()
    document.add_paragraph("1. 測試題幹")
    document.add_paragraph("答案：正確")
    document.save(str(answers_path))
    export_id = await _make_export(answer_docx_path=str(answers_path))

    response = client.get(f"/v1/exports/{export_id}/answers.docx")

    assert response.status_code == 200
    assert "答案卷" in unquote(response.headers["content-disposition"])

    downloaded_path = tmp_path / "downloaded-answers.docx"
    downloaded_path.write_bytes(response.content)
    reopened = new_docx(str(downloaded_path))
    assert [p.text for p in reopened.paragraphs] == ["1. 測試題幹", "答案：正確"]
